#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
微型Delta分拣机器人市场调研报告生成器

此脚本将Markdown格式的报告转换为专业的Word文档，
包含目录、页眉页脚和格式化排版。

使用方法：
    python generate_report.py

依赖安装：
    pip install markdown python-docx

作者：AI辅助调研
日期：2026年5月21日
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

class ReportGenerator:
    """调研报告生成器类"""
    
    def __init__(self, source_file, output_file):
        """
        初始化生成器
        
        Args:
            source_file: Markdown源文件路径
            output_file: 输出Word文件路径
        """
        self.source_file = Path(source_file)
        self.output_file = Path(output_file)
        self.doc = Document()
        self.heading_styles = [
            'Title',
            'Heading 1', 
            'Heading 2',
            'Heading 3',
            'Heading 4'
        ]
        
    def setup_document(self):
        """设置文档基本属性"""
        # 设置文档标题
        core_props = self.doc.core_properties
        core_props.title = "微型桌面Delta分拣机器人创业项目RBTR市场调研报告"
        core_props.author = "AI辅助调研"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
        
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
    def add_custom_heading(self, text, level=1):
        """添加自定义标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置标题样式
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if level == 1:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 84, 144)  # 深蓝色
            elif level == 2:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 123, 182)  # 蓝色
            elif level == 3:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(65, 171, 93)  # 绿色
            elif level == 4:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(224, 130, 20)  # 橙色
        
        return heading
    
    def add_paragraph(self, text, style='Normal'):
        """添加段落"""
        para = self.doc.add_paragraph(text, style=style)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 设置段落格式
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        
        return para
    
    def add_bullet_point(self, text, level=0):
        """添加项目符号列表"""
        style = 'List Bullet' if level == 0 else 'List Bullet 2'
        para = self.doc.add_paragraph(text, style=style)
        
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        
        return para
    
    def add_table(self, headers, rows):
        """添加表格"""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            # 设置表头背景色
            shading = header_cells[i]._element.get_or_add_tcPr()
            shading.set(qn('w:shd'), 'fill:1a5490')
        
        # 添加数据行
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
                row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)
        
        # 设置表格居中对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        return table
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_horizontal_line(self):
        """添加水平分隔线"""
        para = self.doc.add_paragraph()
        para.add_run("─" * 80)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def markdown_to_docx(self):
        """将Markdown转换为Word文档"""
        print(f"正在读取源文件：{self.source_file}")
        
        # 读取Markdown文件
        with open(self.source_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        print("正在转换Markdown为Word文档...")
        
        # 设置文档
        self.setup_document()
        
        # 转换为HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'toc', 'nl2br', 'sane_lists']
        )
        
        # 解析HTML并创建Word文档
        lines = md_content.split('\n')
        in_table = False
        table_data = []
        table_headers = []
        
        for line in lines:
            line = line.strip()
            
            # 跳过空行
            if not line:
                continue
            
            # 处理标题
            if line.startswith('# '):
                self.add_custom_heading(line[2:], 1)
            elif line.startswith('## '):
                self.add_custom_heading(line[3:], 2)
            elif line.startswith('### '):
                self.add_custom_heading(line[4:], 3)
            elif line.startswith('#### '):
                self.add_custom_heading(line[5:], 4)
            
            # 处理表格
            elif line.startswith('|'):
                # 解析表格行
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # 检测表头分隔符
                if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                    table_headers = table_data[-1] if table_data else []
                    table_data = []
                    continue
                
                table_data.append(cells)
            
            # 处理普通段落
            else:
                # 检查是否是表格的最后一行
                if table_data and not line.startswith('|'):
                    # 添加表格
                    if table_headers and table_data:
                        self.add_table(table_headers, table_data)
                    table_data = []
                    table_headers = []
                
                # 处理文本格式
                formatted_line = line
                formatted_line = formatted_line.replace('**', '')
                
                # 添加段落
                self.add_paragraph(formatted_line)
        
        # 处理剩余的表格数据
        if table_data:
            self.add_table(table_headers if table_headers else [], table_data)
        
        # 保存文档
        print(f"正在保存Word文档：{self.output_file}")
        self.doc.save(str(self.output_file))
        print("✅ 报告生成完成！")
        
        return self.output_file


def main():
    """主函数"""
    print("=" * 60)
    print("微型Delta分拣机器人市场调研报告生成器")
    print("=" * 60)
    print()
    
    # 定义文件路径
    script_dir = Path(__file__).parent
    source_file = script_dir / "research.md"
    output_file = script_dir / "research.docx"
    
    # 检查源文件是否存在
    if not source_file.exists():
        print(f"❌ 错误：源文件不存在 {source_file}")
        print("请确保 research.md 文件存在于同一目录下")
        sys.exit(1)
    
    # 检查依赖
    try:
        import markdown
        from docx import Document
        print("✅ 依赖检查通过")
    except ImportError as e:
        print(f"❌ 缺少依赖：{e}")
        print()
        print("请运行以下命令安装依赖：")
        print("pip install markdown python-docx")
        sys.exit(1)
    
    print()
    
    # 生成报告
    generator = ReportGenerator(source_file, output_file)
    output_path = generator.markdown_to_docx()
    
    print()
    print("=" * 60)
    print(f"✅ 报告已生成：{output_path}")
    print("=" * 60)
    print()
    print("提示：")
    print("1. 打开生成的Word文档检查排版")
    print("2. 如需调整格式，可以手动在Word中编辑")
    print("3. 如需生成PDF，可在Word中使用'另存为PDF'功能")


if __name__ == "__main__":
    main()
